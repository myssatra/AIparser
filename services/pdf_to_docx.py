# services/pdf_to_docx.py
import ocrmypdf
from pdf2docx import Converter
from pathlib import Path
import PyPDF2
from docx import Document
from .common import check_tesseract, check_ghostscript

def is_mute_pdf(pdf_path):
    """Проверяет, является ли PDF 'немым' (без встроенного текста) на первых 3 страницах."""
    try:
        with open(pdf_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            for page in pdf.pages[:3]:
                text = page.extract_text()
                if text and text.strip():
                    return False
            return True
    except Exception:
        return True

def process_pdf_to_docx(input_pdfs):
    """Преобразует список PDF-файлов в DOCX с текстом и таблицами, с OCR для 'немых' PDF."""
    # Проверяем зависимости
    tesseract_ok, tesseract_error = check_tesseract()
    if not tesseract_ok:
        return [], tesseract_error
    ghostscript_ok, ghostscript_error = check_ghostscript()
    if not ghostscript_ok:
        return [], ghostscript_error

    if not input_pdfs:
        return [], "Ошибка: выберите хотя бы один PDF-файл."

    output_files = []
    statuses = []

    for input_pdf in input_pdfs:
        input_path = Path(input_pdf)
        temp_ocr_path = input_path.with_stem(f"{input_path.stem}_temp_ocr")
        output_path = input_path.with_suffix(".docx")
        file_name = input_path.name

        try:
            # Проверяем, "немой" ли PDF
            if is_mute_pdf(input_path):
                # Выполняем OCR для "немого" PDF
                ocrmypdf.ocr(
                    input_file=input_path,
                    output_file=temp_ocr_path,
                    language='rus+eng',
                    force_ocr=True,
                    deskew=True,
                    clean=False
                )
                pdf_to_process = temp_ocr_path
            else:
                # Используем исходный PDF для не-"немых"
                pdf_to_process = input_path

            # Извлекаем текст и таблицы с помощью pdf2docx
            cv = Converter(pdf_to_process)
            text = ""
            with open(pdf_to_process, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            tables = cv.extract_tables(start=0, end=len(pdf.pages))
            cv.close()

            # Создаём DOCX с текстом и таблицами
            doc = Document()

            # Добавляем текст (разбиваем на абзацы)
            for paragraph in text.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # Добавляем таблицы
            for table_data in tables:
                if table_data:  # Проверяем, что таблица не пуста
                    rows = len(table_data)
                    cols = max(len(row) for row in table_data) if rows > 0 else 0
                    if rows == 0 or cols == 0:
                        continue

                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            if j < cols:
                                table.cell(i, j).text = str(cell) if cell is not None else ""

            doc.save(output_path)
            output_files.append(str(output_path))
            statuses.append(f"{file_name}: Преобразовано успешно!")

            # Удаляем временный OCR-файл, если он создавался
            if pdf_to_process == temp_ocr_path and temp_ocr_path.exists():
                temp_ocr_path.unlink()

        except Exception as e:
            statuses.append(f"{file_name}: Ошибка: {str(e)}")

    # Объединяем статусы в одну строку с переносами
    status_text = "\n".join(statuses)
    return output_files, status_text