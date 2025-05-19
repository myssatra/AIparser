# services/audio_transcription.py
import whisperx
import torch
from pyannote.audio import Pipeline
import soundfile as sf
import os
import tempfile
import datetime
import time
import ffmpeg
import pandas as pd
import numpy as np
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Включаем TF32 для производительности
torch.backends.cuda.matmul.allow_tf32 = True
torch.backends.cudnn.allow_tf32 = True

# Регистрируем шрифт DejaVu Sans для поддержки кириллицы
try:
    pdfmetrics.registerFont(TTFont('DejaVuSans', './utils/fonts/DejaVuSans.ttf'))
except Exception as e:
    print(f"Ошибка загрузки шрифта DejaVuSans: {str(e)}. Убедитесь, что файл DejaVuSans.ttf доступен.")

def seconds_to_minutes(seconds):
    """Преобразует секунды в минуты и секунды."""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes}:{secs:02d}"

def get_audio_duration(audio_path):
    """Возвращает длительность аудиофайла в секундах."""
    with sf.SoundFile(audio_path) as f:
        duration = len(f) / f.samplerate
    return duration

def extract_audio_from_video(video_path, temp_audio_path):
    """Извлекает аудио из видео и сохраняет во временный файл с помощью FFmpeg."""
    try:
        stream = ffmpeg.input(video_path)
        stream = ffmpeg.output(stream, temp_audio_path, format='wav', acodec='pcm_s16le')
        ffmpeg.run(stream, overwrite_output=True)
        return f"[Лог] Аудио извлечено из видео за {time.time() - start_time:.2f} сек"
    except ffmpeg.Error as e:
        return f"[Лог] Ошибка извлечения аудио: {str(e)}"

def custom_assign_word_speakers(diarization, transcription_result):
    """Кастомная функция для сопоставления спикеров и транскрипции."""
    # Преобразуем диаризацию в DataFrame
    diarize_data = []
    for segment, track, speaker in diarization.itertracks(yield_label=True):
        diarize_data.append({
            'start': segment.start,
            'end': segment.end,
            'speaker': speaker
        })
    diarize_df = pd.DataFrame(diarize_data)
    yield f"[Лог] Столбцы diarize_df: {diarize_df.columns.tolist()}"
    yield f"[Лог] Первые 5 строк diarize_df:\n{diarize_df.head().to_string()}"

    # Копия результата транскрипции
    result = transcription_result.copy()
    for seg in result['segments']:
        seg['speaker'] = 'Неизвестный'  # По умолчанию
        start, end = seg['start'], seg['end']
        # Найти пересекающиеся сегменты диаризации
        overlaps = diarize_df[
            (diarize_df['start'] <= end) & (diarize_df['end'] >= start)
        ]
        if not overlaps.empty:
            # Выбрать спикера с максимальным пересечением
            overlaps['intersection'] = np.minimum(overlaps['end'], end) - np.maximum(overlaps['start'], start)
            overlaps = overlaps[overlaps['intersection'] > 0.1]  # Порог пересечения 0.1 сек
            if not overlaps.empty:
                max_overlap = overlaps['intersection'].idxmax()
                seg['speaker'] = overlaps.loc[max_overlap, 'speaker']
    return result

def transcribe_audio(file_path, model_size="base", output_format="txt", use_diarization=True):
    """Транскрибирует аудио или видео файл с диаризацией и возвращает текст и временный файл."""
    global start_time
    start_time = time.time()
    yield f"[Лог] Начало обработки: {datetime.datetime.now()}"

    if not file_path:
        yield "Ошибка: выберите файл."
        return

    temp_path = file_path
    is_video = file_path.lower().endswith(('.mp4', '.mov', '.mkv', '.avi'))
    if is_video:
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
            temp_path = temp_audio.name
            yield f"[Лог] Извлечение аудио из видео..."
            log = extract_audio_from_video(file_path, temp_path)
            yield log
            if "Ошибка" in log:
                return

    try:
        file_size = os.path.getsize(temp_path)
        max_input_size = 500 * 1024 * 1024  # Лимит входного файла 500 МБ по ТЗ
        warning_threshold = 400 * 1024 * 1024  # Предупреждение при 400 МБ
        if file_size > max_input_size:
            yield f"Ошибка: Размер файла ({file_size / (1024 * 1024):.2f} МБ) превышает лимит 500 МБ."
            return
        elif file_size > warning_threshold:
            yield f"Предупреждение: Размер файла ({file_size / (1024 * 1024):.2f} МБ) близок к лимиту 500 МБ. Возможны проблемы с обработки."

        duration = get_audio_duration(temp_path)
        max_duration = 3 * 60 * 60
        if duration > max_duration:
            yield f"Ошибка: Длительность аудио/видео ({duration / 60:.2f} минут) превышает лимит 3 часов."
            return

        # Устройство как строка
        device = "cuda" if torch.cuda.is_available() else "cpu"
        if device == "cuda":
            torch.cuda.set_device(0)
            yield f"[Лог] GPU устройство установлено: {torch.cuda.get_device_name(0)}"
        else:
            yield f"[Лог] CUDA недоступна, используется CPU"
        compute_type = "float16" if device == "cuda" else "float32"
        yield f"[Лог] Используемое устройство: {device}, тип вычислений: {compute_type}"

        # Загрузка модели Whisper
        yield f"[Лог] Загрузка модели Whisper ({model_size})..."
        model_load_start = time.time()
        model = whisperx.load_model(model_size, device=device, compute_type=compute_type)
        yield f"[Лог] Модель Whisper загружена за {time.time() - model_load_start:.2f} сек"

        # Загрузка аудио
        yield f"[Лог] Загрузка аудиофайла..."
        audio_load_start = time.time()
        audio = whisperx.load_audio(temp_path)
        yield f"[Лог] Аудиофайл загружен за {time.time() - audio_load_start:.2f} сек"

        # Транскрипция
        yield f"[Лог] Начало транскрипции..."
        transcribe_start = time.time()
        result = model.transcribe(audio, language="ru", batch_size=32)
        yield f"[Лог] Транскрипция завершена за {time.time() - transcribe_start:.2f} сек"

        # Диаризация
        diarization = None
        if use_diarization:
            try:
                hf_token = os.getenv("HF_TOKEN")
                yield f"[Лог] Загрузка модели диаризации..."
                diarization_start = time.time()
                diarization_pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1",
                    use_auth_token=hf_token
                )
                if diarization_pipeline is None:
                    yield "[Лог] Не удалось загрузить модель диаризации. Выполняется только транскрипция."
                else:
                    diarization_pipeline.to(torch.device(device))  # Переводим на устройство
                    diarization = diarization_pipeline(temp_path)
                    yield f"[Лог] Результат диаризации: {diarization}"
                    if not diarization:
                        yield "[Лог] Диаризация вернула пустой результат. Пропускаем сопоставление спикеров."
                        diarization = None
                    yield f"[Лог] Диаризация завершена за {time.time() - diarization_start:.2f} сек"
            except Exception as e:
                yield f"[Лог] Ошибка диаризации: {str(e)}"
                diarization = None

        # Форматирование текста
        formatted_text = []
        if diarization is not None:
            try:
                assign_result = custom_assign_word_speakers(diarization, result)
                for item in assign_result:
                    if isinstance(item, str):
                        yield item  # Отладочные сообщения
                    else:
                        result = item  # Финальный результат

                # Объединяем реплики одного спикера
                current_speaker = None
                speaker_text = []
                for segment in result["segments"]:
                    speaker = segment.get("speaker", "Неизвестный")
                    text = segment.get("text", "").strip()
                    if not text:
                        continue

                    # Преобразуем SPEAKER_XX в Спикер N
                    if speaker.startswith("SPEAKER_"):
                        speaker_num = int(speaker.replace("SPEAKER_", "")) + 1
                        speaker_label = f"Спикер {speaker_num}"
                    else:
                        speaker_label = speaker

                    if speaker_label != current_speaker:
                        # Если спикер сменился, сохраняем предыдущий блок
                        if current_speaker is not None:
                            block = f"{current_speaker}: {' '.join(speaker_text)}\n\n"
                            formatted_text.append(block)
                            yield block
                        current_speaker = speaker_label
                        speaker_text = [text]
                    else:
                        # Продолжаем добавлять текст к текущему спикеру
                        speaker_text.append(text)

                # Сохраняем последний блок
                if current_speaker is not None and speaker_text:
                    block = f"{current_speaker}: {' '.join(speaker_text)}\n\n"
                    formatted_text.append(block)
                    yield block

            except Exception as e:
                yield f"[Лог] Ошибка форматирования с диаризацией: {str(e)}"
                import traceback
                yield f"[Лог] Полный стек ошибки: {traceback.format_exc()}"
                # Отладка diarize_df
                try:
                    diarize_data = [
                        {'start': segment.start, 'end': segment.end, 'speaker': speaker}
                        for segment, _, speaker in diarization.itertracks(yield_label=True)
                    ]
                    diarize_df = pd.DataFrame(diarize_data)
                    yield f"[Лог] Столбцы diarize_df: {diarize_df.columns.tolist()}"
                    yield f"[Лог] Первые 5 строк diarize_df:\n{diarize_df.head().to_string()}"
                except Exception as df_error:
                    yield f"[Лог] Ошибка при формировании diarize_df: {str(df_error)}"
                # Продолжить без диаризации
                for segment in result["segments"]:
                    text = segment.get("text", "").strip()
                    if text:
                        line = f"{text}\n"
                        formatted_text.append(line)
                        yield line
        else:
            # Без диаризации просто объединяем весь текст
            for segment in result["segments"]:
                text = segment.get("text", "").strip()
                if text:
                    line = f"{text}\n"
                    formatted_text.append(line)
                    yield line

        output_text = "".join(formatted_text)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Сохранение во временный файл
        try:
            if output_format == "txt":
                with tempfile.NamedTemporaryFile(suffix=f"_transcription_{timestamp}.txt", delete=False, mode='w', encoding='utf-8') as temp_file:
                    temp_file.write(output_text)
                    output_file = temp_file.name
            elif output_format == "docx":
                with tempfile.NamedTemporaryFile(suffix=f"_transcription_{timestamp}.docx", delete=False) as temp_file:
                    output_file = temp_file.name
                    doc = Document()
                    doc.add_heading("Транскрипция аудио/видео", level=1)
                    for line in formatted_text:
                        doc.add_paragraph(line)
                    doc.save(output_file)
            elif output_format == "pdf":
                with tempfile.NamedTemporaryFile(suffix=f"_transcription_{timestamp}.pdf", delete=False) as temp_file:
                    output_file = temp_file.name
                    doc = SimpleDocTemplate(output_file, pagesize=letter)
                    styles = getSampleStyleSheet()
                    styles['Heading1'].fontName = 'DejaVuSans'
                    styles['BodyText'].fontName = 'DejaVuSans'
                    story = [Paragraph("Транскрипция аудио/видео", styles['Heading1'])]
                    for line in formatted_text:
                        story.append(Paragraph(line, styles['BodyText']))
                    doc.build(story)
            else:
                output_file = None
        except Exception as e:
            yield f"[Лог] Ошибка сохранения временного файла: {str(e)}"
            output_file = None

        yield f"[Лог] Обработка завершена за {time.time() - start_time:.2f} сек"
        yield output_text, output_file

    except Exception as e:
        yield f"Ошибка: {str(e)}", None

    finally:
        if is_video and os.path.exists(temp_path):
            os.remove(temp_path)